from docx import Document
import os
import logging
from datetime import datetime
from docx2pdf import convert

class DocumentHandler:
    def __init__(self):
        self.template_path = "attached_assets/chat.docx"
        self.logger = logging.getLogger(__name__)

    def generate_document(self, data, output_format='docx'):
        """Generate a document from template using provided data."""
        try:
            doc = Document(self.template_path)

            for paragraph in doc.paragraphs:
                if not paragraph.text:
                    continue

                text = paragraph.text
                runs = paragraph.runs

                # Zachowaj formatowanie pierwszego runa (jeśli istnieje)
                original_font = None
                original_size = None
                original_bold = None
                original_italic = None
                if runs:
                    original_font = runs[0].font.name
                    original_size = runs[0].font.size
                    original_bold = runs[0].font.bold
                    original_italic = runs[0].font.italic

                if text.startswith("Dla:"):
                    new_text = f"Dla: {data.get('client_name', '')}"
                elif text.startswith("Adres:"):
                    new_text = f"Adres: {data.get('street', '')}, {data.get('postal_code', '')} {data.get('city', '')}"
                elif text.startswith("E-mail:"):
                    new_text = f"E-mail: {data.get('email', '')}"
                elif text.startswith("Tel."):
                    phone = data.get('phone', '').replace('-', ' ')
                    new_text = f"Tel. {phone}"
                elif text.startswith("PANASONIC"):
                    pump_type = "all-in-one" if data.get("all_in_one") == "tak" else "split"
                    kit_text = ""
                    if data.get('kit_number'):
                        kit_text = f"zestaw {data.get('kit_number').upper()}"
                        if data.get('wifi_adapter') == "on":
                            kit_text += " + adapter WiFi"
                    new_text = f"PANASONIC {data.get('pump_model', '')} {data.get('power', '')}kW ({pump_type}) – {kit_text}"
                elif text.startswith("zbiornik ciepłej wody:"):
                    new_text = f"zbiornik ciepłej wody: {data.get('water_tank', '')}"
                elif text.startswith("bufor ciepła:"):
                    heat_buffer = data.get('heat_buffer', '')
                    if not heat_buffer:
                        continue  # Skip this paragraph if heat_buffer is empty
                    new_text = f"bufor ciepła: {heat_buffer}"
                elif "Cena całkowita:" in text:
                    new_text = f"Cena całkowita: **{data.get('price_brutto', '')} zł brutto** (z 8% VAT)"
                else:
                    continue

                # Usuń istniejące runy
                for _ in range(len(paragraph.runs)):
                    paragraph._p.remove(paragraph.runs[0]._r)

                # Dodaj nowy run z zachowaniem oryginalnego formatowania
                run = paragraph.add_run(new_text)
                if original_font:
                    run.font.name = original_font
                if original_size:
                    run.font.size = original_size
                if original_bold is not None:
                    run.font.bold = original_bold
                if original_italic is not None:
                    run.font.italic = original_italic

            # Generate unique filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            os.makedirs("generated_docs", exist_ok=True)
            docx_filename = f"generated_docs/oferta_{timestamp}.docx"
            doc.save(docx_filename)

            return docx_filename

        except Exception as e:
            self.logger.error(f"Error generating document: {str(e)}")
            raise

    def generate_preview(self, data):
        """Generate HTML preview of the document."""
        try:
            preview_template = """
            <div class="preview-document">
                <div class="preview-header">
                    <h1>Oferta dla: {client_name}</h1>
                    <p>Adres: {street}, {postal_code} {city}</p>
                    <p>E-mail: {email}</p>
                    <p>Tel. {phone}</p>
                </div>
                <div class="preview-content">
                    <p><strong>Szczegóły techniczne:</strong><br>
                    PANASONIC {pump_model} {power}kW ({pump_type})<br>
                    Zbiornik CWU: {water_tank}<br>
                    Bufor ciepła: {heat_buffer}</p>
                    <p><strong>Cena całkowita:</strong> {price_brutto} zł brutto (z 8% VAT)</p>
                </div>
            </div>
            """

            preview_data = {
                'client_name': data.get('client_name', ''),
                'street': data.get('street', ''),
                'city': data.get('city', ''),
                'postal_code': data.get('postal_code', ''),
                'email': data.get('email', ''),
                'phone': data.get('phone', ''),
                'pump_model': data.get('pump_model', ''),
                'power': data.get('power', ''),
                'pump_type': 'all-in-one' if data.get('all_in_one') == 'tak' else 'split',
                'water_tank': data.get('water_tank', ''),
                'heat_buffer': data.get('heat_buffer', ''),
                'price_brutto': data.get('price_brutto', '')
            }

            return preview_template.format(**preview_data)

        except Exception as e:
            self.logger.error(f"Error generating preview: {str(e)}")
            raise