from docx import Document
import logging

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def print_document_content():
    try:
        doc = Document("attached_assets/chat.docx")
        logger.info("=== Paragraphs content ===")
        for para in doc.paragraphs:
            logger.info(f"Paragraph text: {para.text}")
            
        logger.info("\n=== Tables content ===")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        logger.info(f"Table cell text: {paragraph.text}")

    except Exception as e:
        logger.error(f"Error reading template: {str(e)}")

if __name__ == "__main__":
    print_document_content()
