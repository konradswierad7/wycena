from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_template():
    doc = Document()
    
    # Add title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run("Oferta dla: {client_name}")
    title_run.bold = True
    title_run.font.size = Pt(14)
    
    # Add empty line
    doc.add_paragraph()
    
    # Add service details
    details = doc.add_paragraph()
    details_run = details.add_run("Szczegóły usługi: {service_details}")
    details_run.font.size = Pt(12)
    
    # Add empty line
    doc.add_paragraph()
    
    # Add price
    price = doc.add_paragraph()
    price_run = price.add_run("Cena: {price} PLN")
    price_run.font.size = Pt(12)
    
    # Add date
    date = doc.add_paragraph()
    date_run = date.add_run("Data: {date}")
    date_run.font.size = Pt(12)
    
    # Add empty line
    doc.add_paragraph()
    
    # Add footer with company details
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    company = footer.add_run("{company_name}\n")
    company.font.size = Pt(11)
    contact = footer.add_run("{contact_person}\n")
    contact.font.size = Pt(11)
    email = footer.add_run("{email}")
    email.font.size = Pt(11)
    
    # Save the document
    doc.save("template.docx")

if __name__ == "__main__":
    create_template()
